# This file will generate a docx summary of the projects in projectsdb

import os
import re
from docx.oxml.shared import OxmlElement, qn

instPrinted = divisionPrinted = departmentPrinted = instCount = divisionCount = departmentCount = 0

def query(db, sqlQuery):
  ''' Run a SQL query against the project database '''
  db.query(sqlQuery)
  r = db.store_result()
  result = r.fetch_row(maxrows=0, how=1)
  return result

def removeNonAscii(s):
  ''' Remove non-ascii and defined html tags from string '''
  filtered = "".join(c for c in s if ord(c)<127 and ord(c)>31)
  # replace br with newline
  filtered = filtered.replace('<br/>','\n')
  filtered = filtered.replace('<br>','\n')
  filtered = filtered.replace('<li>','\n')
  # remove other html
  filtered = re.sub('<[^<]+?>', '', filtered)
  return filtered

def addMetadata(p, text):
  # colored small text
  r = p.add_run(text)
  rPr = r._r.get_or_add_rPr()
  rStyle = OxmlElement('w:color')
  rStyle.set(qn('w:val'), '888888')
  rPr.insert(0, rStyle)
  rStyle = OxmlElement('w:sz')
  rStyle.set(qn('w:val'), '18')
  rPr.insert(0, rStyle)
  rStyle = OxmlElement('w:szCs')
  rStyle.set(qn('w:val'), '18')
  rPr.insert(1, rStyle)

def printOutput(inst, division, department, doc, db, oldest_date):
  global instPrinted, divisionPrinted, departmentPrinted, instCount, divisionCount, departmentCount  

  # Get all project ids of projects that
  #  - have research outcome after and 
  #  - where researcher is FROM ${inst}, ${division} and ${department} and
  #  - the project used the Pan cluster
  #  - the research output has been collected after the survey in 2013
  tmpquery = '''SELECT DISTINCT project.id FROM project
      INNER JOIN researcher_project rp ON
          rp.projectId=project.id 
      INNER JOIN researcher r ON
          r.id=rp.researcherId AND
          r.institution='%s' AND 
          r.division='%s' AND r.department='%s' 
      INNER JOIN researchoutput ro ON
          project.id=ro.projectId AND
          ro.date > '%s'
      INNER JOIN project_facility pf ON
          project.id=pf.projectId AND
          pf.facilityId=1''' % (inst, division, department, oldest_date)

  projects_with_research_outcome = query(db, tmpquery)
  pids = set([tmp['id'] for tmp in projects_with_research_outcome])

  if not pids:
    return

  outputs=[]

  for pid in pids:

    # get name of project owner
    tmpquery = '''SELECT researcher.fullName as name FROM researcher
        INNER JOIN researcher_project rp ON
            researcher.id=rp.researcherId AND
            rp.projectId=%s AND
            rp.researcherRoleId = 1''' % pid 
    researcher = query(db, tmpquery)[0]['name']

    # check if the owner of the project is affiliated with the currently processed department
    affil = query(db, 'SELECT fullName, institution, division, department from researcher \
                     INNER JOIN researcher_project rp ON researcher.id = rp.researcherId AND rp.projectId=%s \
                     INNER JOIN researcherrole rr ON rr.id = rp.researcherRoleId AND rr.name=\'Project Owner\'' % pid)[0]
    
    if affil['institution'] == inst and affil['division'] == division and affil['department'] == department:
      # get ids of research output ordered by output type
      outputIds = query(db, '''SELECT id FROM researchoutput WHERE projectId=%s AND date > '%s' ORDER BY typeId''' % (pid, oldest_date))
      if outputIds:
        # print institution if not already printed
        if instPrinted == 0:
          instCount = instCount + 1
          doc.add_heading('%s' % inst, 2)
          # insert graph if available 
          graph_pic = 'pics/ResearchOutputs_%s.png' % inst.replace(' ','-')
          if os.path.exists(graph_pic):
            doc.add_picture(graph_pic)
          instPrinted = 1
        # print division if not already printed
        if divisionPrinted == 0 and division != '':
          divisionCount = divisionCount + 1
          doc.add_heading('%s' % division, 3)
          # insert graph if available 
          graph_pic = 'pics/ResearchOutputs_%s_%s.png' % (inst.replace(' ','-'), division.replace(' ','-'))
          if os.path.exists(graph_pic):
            doc.add_picture(graph_pic)
          divisionPrinted = 1
        # print department if not already printed
        if departmentPrinted == 0 and department != '':
          departmentCount = departmentCount + 1
          doc.add_heading('%s' % department, 4)
          departmentPrinted = 1

        for outputid in [tmp['id'] for tmp in outputIds]:
          typeName = query(db, 'SELECT name FROM researchoutputtype WHERE id=(SELECT typeId FROM researchoutput WHERE id=%s)' % outputid)[0]['name']
          desc = query(db, 'SELECT description FROM researchoutput WHERE id=%s' % outputid)[0]['description']
          pcode = query(db, 'SELECT projectCode FROM project WHERE id=%s' % pid)[0]['projectCode']
          p = doc.add_paragraph('')
          p.add_run('%s' % removeNonAscii(desc)).add_break()
          addMetadata(p, '[Type = %s | ProjectOwner = %s | ProjectCode = %s]' % (typeName, researcher, pcode))


def add_research_output(db, doc, institution, oldest_date):
  global divisionPrinted, departmentPrinted, departmentCount  
  doc.add_heading('Research Output', 1)
  txt = 'This section lists research output that was generated with the help of the NeSI Auckland cluster. '
  txt += 'These results were collected between 2013-07-05 and today and are sorted by institution, division and department.'
  doc.add_paragraph(txt) 
  txt = 'Overview diagrams that illustrate the number and ratio of research outcomes are provided on the '
  txt += 'institution and the division level, and citations of each research outcome are provided on the '
  txt += 'department level.'
  doc.add_paragraph(txt) 
  txt = 'The research outcomes have been collected through individual consultations with researchers, and in '
  txt += 'the annual Centre for eResearch survey.'
  doc.add_paragraph(txt) 

  divisions = query(db, '''SELECT DISTINCT division FROM researcher WHERE institution='%s' ORDER BY division''' % institution)
  for division in [tmp['division'] for tmp in divisions]:
    divisionPrinted = 0
    departments = query(db, '''SELECT DISTINCT department FROM researcher WHERE institution='%s' AND division=\'%s\' ORDER BY department''' % (institution, division))
    for department in [tmp['department'] for tmp in departments]:
      departmentPrinted = 0
      printOutput(institution, division, department, doc, db, oldest_date)
    departmentCount = 0
  
