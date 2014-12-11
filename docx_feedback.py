import re
from docx.oxml.shared import OxmlElement, qn

instPrinted = divisionPrinted = departmentPrinted = instCount = divisionCount = departmentCount = 0

def query(db, sqlQuery):
  ''' Run a SQL query against the project database '''
  db.query(sqlQuery)
  r = db.store_result()
  result = r.fetch_row(maxrows=0, how=1)
  return result

def removeNonAscii(s):
  ''' Remove non-ascii and defined html tags from string '''
  filtered = "".join(c for c in s if ord(c)<127 and ord(c)>31)
  # replace br with newline
  filtered = filtered.replace('<br/>','\n')
  filtered = filtered.replace('<br>','\n')
  filtered = filtered.replace('<li>','\n')
  # remove other html
  filtered = re.sub('<[^<]+?>', '', filtered)
  return filtered

def print_quote(doc, text):
  p = doc.add_paragraph('')
  r = p.add_run(text)
  r.italic = True
  pPr = p._p.get_or_add_pPr()
  pStyle = OxmlElement('w:ind')
  pStyle.set(qn('w:left'), '1000')
  pStyle.set(qn('w:right'), '0')
  pPr.insert(0, pStyle)
  pStyle = OxmlElement('w:jc')
  pStyle.set(qn('w:val'), 'both')
  pPr.insert(1, pStyle)

def print_metadata(doc, text):
  p = doc.add_paragraph('')
  r = p.add_run(text)
  rPr = r._r.get_or_add_rPr()
  rStyle = OxmlElement('w:color')
  rStyle.set(qn('w:val'), '888888')
  rPr.insert(0, rStyle)
  rStyle = OxmlElement('w:sz')
  rStyle.set(qn('w:val'), '18')
  rPr.insert(1, rStyle)
  rStyle = OxmlElement('w:szCs')
  rStyle.set(qn('w:val'), '18')
  rPr.insert(2, rStyle)

def printFeedback(inst, division, department, doc, db, oldest_date):
  global instPrinted, divisionPrinted, departmentPrinted, instCount, divisionCount, departmentCount
  
  tmpquery = '''SELECT pf.id from projectfollowup pf 
      INNER JOIN researcher r ON
          r.id=pf.researcherId AND
          r.institution='%s' AND 
          r.division='%s' AND 
          r.department='%s' 
      WHERE pf.date > '%s';''' % (inst, division, department, oldest_date)

  feedbackIds = query(db, tmpquery)
  fbIds = set([tmp['id'] for tmp in feedbackIds])
  outputs=[]

  if not fbIds:
    return

  for fbId in fbIds:
    # get name of project owner
    tmpquery = '''SELECT fullName, institution, department, division, ir.name as role FROM researcher r
        INNER JOIN institutionalrole ir ON
            r.institutionalRoleId = ir.id
        INNER JOIN projectfollowup pf ON
            r.id = pf.researcherId 
        WHERE pf.id=%s''' % fbId 
    researcher = query(db, tmpquery)[0]
    rName = researcher['fullName']
    rInst = researcher['institution']
    rDiv = researcher['division']
    rDep = researcher['department']
    rRole = researcher['role']

    if rInst == inst and rDiv == division and rDep == department:
      pcode = query(db, '''SELECT projectCode FROM project p
          INNER JOIN projectfollowup pf ON
            pf.projectId = p.id
          WHERE pf.id=%s''' % fbId)[0]['projectCode']
      fb = query(db, '''SELECT SUBSTRING(pf.notes, LOCATE('Feedback:<br>', pf.notes)) AS feedback
          FROM projectfollowup pf 
          WHERE id = %s''' % fbId)[0]['feedback']
      fb = fb.replace('Feedback:<br>','')

      if fb.strip() and fb.strip() != 'N/A':
        # print institution if not already printed
        if instPrinted == 0:
          instCount = instCount + 1
          doc.add_heading('%s' % inst, 2)
          instPrinted = 1
        # print division if not already printed
        if divisionPrinted == 0 and division:
          divisionCount = divisionCount + 1
          doc.add_heading('%s' % division, 3)
          divisionPrinted = 1
        # print department if not already printed
        if departmentPrinted == 0 and department:
          departmentCount = departmentCount + 1
          doc.add_heading('%s' % department, 4)
          departmentPrinted = 1

        #p = doc.add_paragraph('%s (%s):' % (rName,rRole))
        doc.add_paragraph('%s:' % rRole)
        print_quote(doc, removeNonAscii(fb.strip())) 
        print_metadata(doc, '[ProjectCode = %s]' % pcode)
        doc.add_paragraph(' ')

# main 

def add_feedback(db, doc, institution, oldest_date):
  global divisionPrinted, departmentPrinted, departmentCount
  doc.add_heading('Feedback from researchers', 1)
  txt = 'As part of the Centre for eResearch annual survey conducted on 2014-07-07 researchers '
  txt += 'had the opportunity to give feedback on the service provided by the Centre for eResearch.'
  doc.add_paragraph(txt)
  txt = 'The following feedback has been reported in replies to this survey, sorted by '
  txt += 'institution, division and department.'
  doc.add_paragraph(txt)

  institutions = query(db, 'SELECT DISTINCT institution FROM researcher ORDER BY institution')

  divisions = query(db, 'SELECT DISTINCT division FROM researcher WHERE institution=\'%s\' ORDER BY division' % institution)
  for division in [tmp['division'] for tmp in divisions]:
    divisionPrinted = 0
    departments = query(db, '''SELECT DISTINCT department FROM researcher WHERE institution='%s' AND division='%s' ORDER BY department''' % (institution, division))
    for department in [tmp['department'] for tmp in departments]:
      departmentPrinted = 0
      printFeedback(institution, division, department, doc, db, oldest_date)
    departmentCount = 0
  divisionCount = 0
  

