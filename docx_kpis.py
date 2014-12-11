import os
import re
import sys
import _mysql
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.section import WD_SECTION, WD_ORIENT

instPrinted = divisionPrinted = departmentPrinted = instCount = divisionCount = departmentCount = 0
orientation = 'portrait'

scaling_categories = { 1: 'Throughput', 2: 'CPUs', 3: 'Memory', 4: 'Optimisation' }

def query(db, sqlQuery):
  ''' Run a SQL query against the project database '''
  db.query(sqlQuery)
  r = db.store_result()
  result = r.fetch_row(maxrows=0, how=1)
  return result

def add_scaling_graphs(doc, institution, division=None):
  table = doc.add_table(rows=2, cols=2, style='TableGrid')
  count = 1
  for i in range(0,2):
    cells = table.rows[i].cells
    for j in range(0,2):
      p = cells[j].paragraphs[0]
      r = p.add_run()
      if division:
        pic = 'pics/KPI_%s_%s_%s' % (institution.replace(' ','-'), division.replace(' ','-'), scaling_categories[count])
      else:
        pic = 'pics/KPI_%s_%s' % (institution.replace(' ','-'), scaling_categories[count])
      pic += '.png'
      if os.path.exists(pic):
        r.add_picture(pic)
      else:
        print '%s does not exist' % pic
      count += 1

def printScaledProjects(inst, division, doc, db, oldest_date):
  global instPrinted, divisionPrinted, departmentPrinted, instCount, divisionCount, departmentCount, scaling_categories, orientation
 
  departments = query(db, 'SELECT DISTINCT department FROM researcher WHERE institution=\'%s\' AND division=\'%s\' ORDER BY department' % (inst, division))

  for department in [tmp['department'] for tmp in departments if tmp['department'].strip()]:
    departmentPrinted = 0
    # get all projects that have a kpi 
    q = '''SELECT DISTINCT projectId FROM project_kpi pk 
           INNER JOIN project p ON p.id = pk.projectId
           WHERE pk.date >= '%s' AND
             p.hostInstitution='%s' AND 
             p.division='%s' AND 
             p.department='%s' ''' % (oldest_date, inst, division, department)

    projectIds = query(db, q)
    pids = [tmp['projectId'] for tmp in projectIds]
    if not pids:
      continue

    # print institution if not already printed
    if not instPrinted:
      instCount = instCount + 1
      doc.add_heading('%s' % inst, 2)
      add_scaling_graphs(doc, inst)
      instPrinted = 1
    # print division if not already printed
    if not divisionPrinted and division:
      divisionCount = divisionCount + 1
      doc.add_heading('%s' % division, 3)
      add_scaling_graphs(doc, inst, division)
      section = doc.add_section(WD_SECTION.NEW_PAGE)
      new_width, new_height = section.page_height, section.page_width
      section.page_width = new_width
      section.page_height = new_height
      section.right_margin = Inches(0.5)
      section.left_margin = Inches(0.5)
      section.top_margin = Inches(0.5)
      section.bottom_margin = Inches(0.5)
      section.orientation = WD_ORIENT.LANDSCAPE
      orientation = 'landscape'
      divisionPrinted = 1
    # print department if not already printed
    if not departmentPrinted and department:
      departmentCount = departmentCount + 1
      doc.add_heading('%s' % department, 4)
      departmentPrinted = 1

    table = doc.add_table(rows=len(pids) + 1, cols=6, style='MediumShading1-Accent2')
    cells = table.rows[0].cells
    cells[0].text = 'Project Title'
    cells[1].text = 'Project Owner'
    cells[2].text = 'Scaling Factor Throughput'
    cells[3].text = 'Scaling Factor CPU'
    cells[4].text = 'Scaling Factor Memory'
    cells[5].text = 'Scaling Factor Optimisation'

    count = 1
    for pid in pids:
      # get name of project owner
      q = '''SELECT code, MAX(value) as value from project_kpi
             WHERE projectId = %s AND date >= '%s'
             GROUP BY code''' % (pid, oldest_date)
      tmp = query(db, q)
      kpis = {}
      for kpi in tmp:
        kpis[int(kpi['code'])] = kpi['value']
      q = '''SELECT DISTINCT p.name AS project_title, p.projectCode AS project_code, r.fullName AS researcher_name FROM project p
             INNER JOIN researcher_project rp ON rp.projectId = p.id
             INNER JOIN researcher r ON rp.researcherId = r.id
             WHERE rp.researcherRoleId = 1 AND
               p.id = %s''' % pid
      res = query(db, q)
      if not res:
        raise Exception('Unexpected empty result set for query')
      cells = table.rows[count].cells
      cells[0].text = res[0]['project_title']
      cells[1].text = res[0]['researcher_name']
      for i in range(1, 5):
        if i in kpis.keys():
          cells[i+1].text = kpis[i]
        else:
          cells[i+1].text = '-'
      count += 1  

  if orientation == 'landscape':
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.right_margin = Inches(1)
    section.left_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.orientation = WD_ORIENT.PORTRAIT

    orientation = 'portrait'

def add_kpis(db, doc, institution, oldest_date): 
  global divisionPrinted, departmentPrinted, departmentCount, divisionCount
  doc.add_heading('Scaling', 1)
  txt = 'In many ways the NeSI Auckland cluster offers significant advantages over desktop computers for compute-intensive projects. '
  txt += 'This is because the cluster helps overcome common limiting factors such as the availability of memory and the number of processors available. '
  txt += 'This section illustrates how the cluster has helped to scale projects across the following limiting factors:'
  doc.add_paragraph(txt)

  table = doc.add_table(rows=4, cols=2, style='MediumList1-Accent2')
  table.rows[0].cells[0].text = 'Throughput'
  table.rows[1].cells[0].text = 'CPUs'
  table.rows[2].cells[0].text = 'Memory'
  table.rows[3].cells[0].text = 'Optimisation'

  txt = 'The cluster enables researchers to run more jobs at a time than they would be able to run without the cluster. '
  txt += 'A typical scenario is where several similar simulations are performed to test the importance of key parameters in an underlying model. '
  txt += 'Being able to test various conditions simultaneously can greatly improve a researcher\'s productivity.'
  table.rows[0].cells[1].text = txt

  txt = 'Many scientific software applications can utilise multiple CPU cores concurrently across one or more computers. '
  txt += 'The cluster enables researchers to run such applications with a much higher CPU core count than they would otherwise have access to. '
  txt += 'This reduces the amount of time required to produce results and allows researchers to expand the scope of their investigations.'
  table.rows[1].cells[1].text = txt

  txt = 'Each cluster node has more memory than a standard desktop computer. '
  txt += 'This category covers scenarios where researchers can work with larger data sets or more detailed models because of the amount '
  txt += 'of memory available to them, either on a single cluster node, or by leveraging memory across multiple cluster nodes.'
  table.rows[2].cells[1].text = txt

  txt = 'This category covers scenarios in which research workflows have been optimised to reduce the amount of time required to generate results. '
  txt += 'This includes general improvements made to research codes as well as extensions to enable parallel processing. '
  txt += 'Improved software configuration and installation is also included in this category.'
  table.rows[3].cells[1].text = txt

  doc.add_paragraph('')
  txt = 'Each of the graphs in the following section shows the total number of research projects that have been scaled (horizontal axis) '
  txt += 'across various scaling bands (vertical axis). '
  txt += 'The scaling factor value shows the improvement in productivity from using the cluster.'
  doc.add_paragraph(txt)
  txt = 'For example, the graph below shows that 2 projects have reported an improvement of between 8-16 times due to the higher number of CPU cores available to them. '
  txt += 'In this case, one project may have a scaling of 10 while the other has a scaling of 15, so they are summarized as having a '
  txt += 'scaling of between 8-16.'
  doc.add_paragraph(txt)
  doc.add_picture('pics/KPI_sample.png')

  doc.add_page_break()

  divisions = query(db, 'SELECT DISTINCT division FROM researcher WHERE institution=\'%s\' ORDER BY division' % institution)
  section = None
  for division in [tmp['division'] for tmp in divisions]:
    divisionPrinted = 0
    printScaledProjects(institution, division, doc, db, oldest_date)
